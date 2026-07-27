# -*- coding: utf-8 -*-
"""
Build the Mular meeting script (.docx).

A speaking document, not a report: short lines, one idea per line, the exact
words in quotes so they can be read off the page under pressure, and the
reasoning kept in the margin notes rather than in the sentences to be spoken.
Mular's navy/green so it sits in the same family as the Jasiri assessment.

    python scripts/build-mular-meeting-script.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x00, 0x3C, 0x71)
GREEN = RGBColor(0x50, 0x95, 0x1D)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x60, 0x60, 0x60)
FAINT = RGBColor(0x8A, 0x8A, 0x8A)
RED = RGBColor(0xB0, 0x2A, 0x2A)

OUT = os.path.join("reports", "Mular-Meeting-Script.docx")


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def rule(doc, color="003C71", size=12):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    b.append(bottom)
    pPr.append(b)
    p.paragraph_format.space_after = Pt(6)
    return p


def para(doc, text="", size=10.5, bold=False, italic=False, color=INK,
         before=0, after=4, indent=0, font="Calibri", align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent:
        pf.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.name = font
    return p


def say(doc, text, indent=0.6):
    """A line to be spoken, verbatim. Set apart so the eye finds it instantly."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(indent)
    pf.space_before = Pt(3)
    pf.space_after = Pt(5)
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "50951D")
    b.append(left)
    pPr.append(b)
    r = p.add_run(u"“" + text + u"”")
    r.font.size = Pt(11)
    r.font.name = "Georgia"
    r.font.color.rgb = INK
    return p


def note(doc, text, indent=0.6):
    """Stage direction — never read aloud."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(indent)
    pf.space_before = Pt(0)
    pf.space_after = Pt(7)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = FAINT
    r.font.name = "Calibri"
    return p


def bullet(doc, text, size=10.5, indent=0.85, color=INK, bold_lead=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(indent)
    pf.space_after = Pt(3)
    r0 = p.add_run(u"•   ")
    r0.font.size = Pt(size)
    r0.font.color.rgb = GREEN
    r0.font.name = "Calibri"
    if bold_lead:
        rb = p.add_run(bold_lead)
        rb.font.size = Pt(size)
        rb.font.bold = True
        rb.font.color.rgb = INK
        rb.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def h1(doc, num, text):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(num)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = GREEN
    r.font.name = "Calibri"
    r2 = p.add_run("   " + text)
    r2.font.size = Pt(17)
    r2.font.bold = True
    r2.font.color.rgb = NAVY
    r2.font.name = "Calibri"
    rule(doc, "DDDDDD", 6)


def h2(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def table(doc, headers, rows, widths=None, head_fill="003C71"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        shade(hdr[i], head_fill)
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(htxt)
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Calibri"
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            r.font.color.rgb = INK
            r.font.name = "Calibri"
            if i == 0:
                r.font.bold = True
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ─────────────────────────────────────────────────────────────────────────────
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)

# ── Cover ────────────────────────────────────────────────────────────────────
para(doc, "MULAR CREDIT LTD", 10, bold=True, color=GREEN, after=2)
para(doc, "Meeting Script", 30, bold=True, color=NAVY, after=0)
para(doc, "What to say, in what order, and in whose words", 13, color=MUTED, after=14)
rule(doc)
para(doc, "Emmanuel Birgen  ·  IT Consultant, Serve Well Co.", 10.5, bold=True, after=1)
para(doc, "Follows the Jasiri LMS Independent Assessment of 24 July 2026", 9.5, color=MUTED, after=1)
para(doc, "27 July 2026", 9.5, color=MUTED, after=14)

para(doc, "HOW TO USE THIS DOCUMENT", 9, bold=True, color=GREEN, after=3)
bullet(doc, "are the words to say. You can read them aloud as written.", bold_lead="Green-barred lines ")
bullet(doc, "are stage directions — what to do, never what to read out.", bold_lead="Grey italics ")
bullet(doc, "Do not read this document front to back in the room. Sections 1–3 are the spine; the rest is there when you need it.")
bullet(doc, "If you only get ten minutes, do Section 1, Section 4, and the ask in Section 8.")

doc.add_paragraph()
para(doc, "THE ONE THING TO GET RIGHT", 9, bold=True, color=RED, after=3)
para(doc, "You are not selling software to people who have software. You are selling TURNAROUND "
          "— the hours between a customer asking and money landing — to people who make "
          "their money on volume. Every technical point in this script exists to serve that sentence. "
          "If a thread stops serving it, drop the thread.", 10.5, after=6)

# ── 1 ────────────────────────────────────────────────────────────────────────
h1(doc, "1", "Opening — the first ninety seconds")

note(doc, "Set the frame before anyone else does. You are the person who read their system "
          "properly, on their own tenant, and wrote it down. That is your standing in the room.")

say(doc, "Thank you for the access. I want to be useful rather than polite, so I will be direct — "
         "and I will start with what is working, because there is a lot of it.")

say(doc, "Jasiri is a capable platform. The analytics are genuinely strong, the product model is "
         "sensible, and lead management is better than most of what I see. Nothing I am going to "
         "say takes that away.")

note(doc, "Say this and mean it. Somebody in that room chose Jasiri, and possibly defended the "
          "choice. If you open by attacking it, everything after that is heard as a sales pitch. "
          "Praising the analytics costs you nothing and buys you the right to be believed later.")

say(doc, "What I found is that Jasiri is being held back by three things: onboarding friction that "
         "stretches your turnaround, a few areas that look finished but are not, and some security "
         "exposures on the live system. None of it is fatal. All of it is fixable. And the same "
         "fixes are the road to a much faster origination engine.")

h2(doc, "If they ask up front what this meeting is for")
say(doc, "Two things. First, to walk you through the findings so you can act on them — whoever "
         "does the work. Second, to show you what I have already built, so you can judge whether "
         "you want me to build it for Mular or licence the parts you need.")

note(doc, "This is the honest frame and it is also the strongest one. It says the report was not "
          "bait. It makes it safe for them to keep Jasiri and still buy from you.")

# ── 2 ────────────────────────────────────────────────────────────────────────
h1(doc, "2", "The headline — turnaround is the money")

note(doc, "Do not lead with AI. Lead with the number that runs their business.")

say(doc, "Your biggest opportunity is not a feature. It is turnaround — how long it takes from a "
         "customer asking to money in their M-Pesa. Everything else is downstream of that.")

say(doc, "Two capabilities collapse that time without loosening a single control.")

say(doc, "The first is identity-first onboarding. An ID number comes back with verified particulars "
         "from IPRS — the exact thing that impressed you when you handed over only your own ID. "
         "The officer stops typing and starts confirming.")

say(doc, "The second is internal statement scoring. We read the customer's M-Pesa statement and it "
         "produces three things at once: a score, a starting limit, and the product they actually "
         "qualify for — each with reason codes an officer can read out loud.")

say(doc, "Put those together and expensive human judgement moves to the few files that need it, "
         "instead of every file.")

h2(doc, "The line to land it")
say(doc, "Right now every customer costs you the same amount of officer time, whether they are "
         "obviously good or obviously marginal. That is the leak. You are paying full price for "
         "decisions that make themselves.")

note(doc, "Pause here. This is the sentence that makes a CEO lean forward. Let them respond before "
          "you carry on.")

# ── 3 ────────────────────────────────────────────────────────────────────────
h1(doc, "3", "The security conversation — handle with care")

note(doc, "This is the most delicate part of the meeting and the easiest to get wrong. "
          "It is leverage ONLY if you never once use it as leverage.")

para(doc, "Three rules", 10.5, bold=True, color=NAVY, before=4, after=3)
bullet(doc, "Say it privately, calmly, and only once. Do not repeat it for effect.", bold_lead="Understate it. ")
bullet(doc, "You held it back, disclosed it in writing, changed nothing, took nothing.", bold_lead="Emphasise restraint. ")
bullet(doc, "Fixing these is not a reason to hire you. Say so out loud.", bold_lead="Separate it from the sale. ")

say(doc, "There is a section in the report I want to flag directly, and I would rather say it to "
         "your face than let you find it on page twelve.")

say(doc, "During normal testing, an Administrator account was able to promote itself to Super Admin. "
         "Sessions also do not expire. I confirmed both, I took nothing, I changed nothing beyond "
         "the minimum needed to prove it, and I put it in writing to you rather than anywhere else.")

say(doc, "I would fix those two this week, whoever does the work for you. It is not a reason to "
         "hire me. It is just the right thing to do.")

note(doc, "That last line is the most valuable sentence in the whole meeting. It converts a finding "
          "that could feel like a threat into a demonstration of exactly what kind of vendor you are. "
          "Do not skip it. Do not soften it into a maybe.")

h2(doc, "If they get defensive")
say(doc, "I am not raising it to embarrass anybody. Every platform I review has something — this "
         "is a normal finding, handled the normal way. The only thing that matters is that it gets "
         "closed.")

# ── 4 ────────────────────────────────────────────────────────────────────────
h1(doc, "4", "The demo — run it in this order")

note(doc, "Do not tour the product. Tell one customer's story end to end, and let them watch it happen. "
          "Whole thing in twelve minutes. If they interrupt with questions, you are winning — stop and answer.")

table(doc,
      ["#", "Show", "The line that goes with it"],
      [
          ["1", "BirgenAI Hub → Loans. Mular is first on the shelf.",
           "“This is a customer pipeline, not a directory. Mular sits at the top of the Loans shelf. "
           "A user installs it and it goes into their My Apps — so you are on their phone, next to their bank.”"],
          ["2", "Tap Mular. It opens Mular's own portal, your branding.",
           "“They never leave your brand. It is your portal, your logo, your colours — we are underneath it.”"],
          ["3", "Onboard as a new customer: phone, code, statement.",
           "“That is a customer who arrived from the Hub and is now a lead in your book. That is the stream.”"],
          ["4", "The crunch: score, starting limit, matched product, reason codes.",
           "“The officer did not decide this — but they can explain every line of it to the customer. "
           "That is the difference between automation and a black box.”"],
          ["5", "Switch to the LMS console. Same customer, same stage.",
           "“Same person, two screens, one status. If she is in KYC on her phone, she is in KYC on your officer's screen. "
           "They cannot drift, because there is only one thing deciding.”"],
          ["6", "Send the offer. Sign it on the phone.",
           "“The agreement goes out from here and she signs it there, on her own handset. "
           "That is the contract — the terms she saw, the moment she accepted, and how we know it was her.”"],
          ["7", "Customer-360 → Reset their portal PIN.",
           "“Returning customers sign in with their national ID and a PIN — no SMS, no waiting, no cost to you. "
           "And when they lose it, your officer fixes it in one click, on the call.”"],
          ["8", "ServiceSuite AI dock — open it, show the three apps.",
           "“Three different intelligences, one at a time, like apps on a phone. "
           "You can never be typing to one and sending to another.”"],
          ["9", "Analytics: ask a live question. Show the SQL. Export the PDF.",
           "“It answers from your live book, it shows you the query it ran, and it hands you the report. "
           "A number you cannot check is a number you cannot act on.”"],
          ["10", "Closed ML Loop module.",
           "“This is why I built the whole thing. I will come to it.”"],
      ],
      widths=[0.9, 5.4, 9.5])

h2(doc, "The two moments that sell it")
bullet(doc, "— when the same customer appears on both screens at the same stage. Say the line, then stop talking and let them look.",
       bold_lead="Step 5 ")
bullet(doc, "— signing on the phone. If you can get someone in the room to sign on their own handset, do it. Nothing you say beats that.",
       bold_lead="Step 6 ")

# ── 5 ────────────────────────────────────────────────────────────────────────
h1(doc, "5", "The pipeline story — where their customers come from")

note(doc, "This is the part Jasiri cannot answer at all, and it is the one with the clearest money attached. "
          "Spend time here.")

say(doc, "Every lender I meet has the same two problems: finding customers, and deciding on them "
         "quickly. Your current system helps with the second one. Nothing helps with the first.")

say(doc, "BirgenAI Hub is a consumer app with a Loans shelf. Mular is the first thing on it. "
         "Somebody browsing for credit sees you before they see anybody else, installs your app, "
         "and lands in your funnel — already identified.")

say(doc, "Those arrive in your console as leads. Not names on a list — people who already "
         "gave consent, already have a phone number we verified, and in many cases already have "
         "a score.")

h2(doc, "If they ask how many users the Hub has")
note(doc, "Do not inflate. A number they can check and disbelieve costs you the whole meeting.")
say(doc, "I will give you the honest answer: it is early, and I would rather show you the mechanism "
         "than quote you a number you cannot verify. What I can commit to is that Mular holds the "
         "first position on that shelf, and that every install lands in your book, not in a "
         "marketplace where you compete on the same screen.")

# ── 6 ────────────────────────────────────────────────────────────────────────
h1(doc, "6", "The data-science story — the part nobody else has")

note(doc, "This is your real differentiator and it is also the one most likely to be oversold by a "
          "competitor. Win it by being the only person in the room who states a limit.")

say(doc, "I am a data scientist. I did not build a lending system because I wanted to be in "
         "software. I built it because a credit model is only as good as the outcomes you can "
         "feed it, and almost nobody is capturing them properly.")

say(doc, "Every decision this platform makes stores the exact inputs behind it. Months later, when "
         "we know whether that loan was repaid or defaulted, the answer is written back onto the "
         "same record. That is a training set being assembled in production, one customer at a time.")

say(doc, "At three hundred completed outcomes, a model trained on Mular's own borrowers takes over "
         "from the general one. From that day, every decision makes the next one better.")

h2(doc, "The honesty that closes it")
say(doc, "And I will tell you what most people will not: below that number, nobody can honestly "
         "claim their model is tuned to your book. Not me, not anyone. What I can show you is "
         "exactly where you are on that road, and the screen that tracks it.")

note(doc, "Then open the Closed ML Loop module. Show the Wilson interval panel — the same default "
          "rate measured at four sample sizes, the estimate barely moving while the uncertainty "
          "collapses. Say:")

say(doc, "This is the argument for volume, and it is arithmetic, not opinion. More customers do "
         "not magically improve the model. They make it possible to TELL whether it improved.")

h2(doc, "The two errors — use this if there is a risk or finance person in the room")
say(doc, "A credit model makes two mistakes and they do not cost the same. Approving someone who "
         "defaults costs you the principal. Declining someone who would have repaid costs you the "
         "margin. In this market the first is roughly seven times worse — so the model is tuned "
         "to catch defaults, and I can show you what that choice costs and earns in shillings.")

# ── 7 ────────────────────────────────────────────────────────────────────────
h1(doc, "7", "Objections — and the words that answer them")

table(doc,
      ["If they say", "Say this"],
      [
          ["“We already have Jasiri.”",
           "“Good — and you may want to keep it. There are three ways to work with me and only one of "
           "them replaces anything. The other two make what you already have faster.”"],
          ["“This is a big change.”",
           "“Then do not start with a big change. Start with the security fixes and IPRS onboarding. "
           "That is weeks, not quarters, and you will feel it in your turnaround immediately.”"],
          ["“How do we know your scoring works?”",
           "“You do not, yet — and you should not take my word for it. Run it in shadow beside your "
           "current process. It scores every application, changes no decision, and in ninety days you "
           "will have your own evidence instead of my claim.”"],
          ["“What happens if you get hit by a bus?”",
           "“Fair question. The engines are APIs with documented contracts, the data is in your own "
           "database, and I will hand over the schema and the runbook on day one. You are not buying "
           "a person, you are buying something that keeps working without me.”"],
          ["“What does it cost?”",
           "“It depends which of the three you want, and I would rather scope it properly than quote "
           "you a number in a meeting. Give me the shape you want and you will have a written proposal "
           "in forty-eight hours.”"],
          ["“Can you fix the Jasiri issues for us?”",
           "“Some of them are in their code, not yours — those are a conversation with your vendor and "
           "I will help you have it. The onboarding and scoring layer I can do without touching Jasiri "
           "at all.”"],
          ["“Is our data safe with an AI?”",
           "“The analytics run read-only against your own database, scoped by the database itself. It "
           "can measure your book; it cannot change a row in it. And every answer shows you the query "
           "it ran.”"],
          ["“Why should we trust a small firm?”",
           "“Because you have already seen how I work. I reviewed your live system, found something "
           "serious, and brought it to you privately with nothing attached to it. That is the same way "
           "I would handle your customers' data.”"],
      ],
      widths=[5.2, 10.6])

# ── 8 ────────────────────────────────────────────────────────────────────────
h1(doc, "8", "The ask — three doors, one decision")

note(doc, "Never leave a meeting with 'we will think about it'. Offer three sizes so the answer is "
          "WHICH, not WHETHER. Present them in this order — smallest first — so the big one "
          "sounds reasonable by the time you reach it.")

table(doc,
      ["Door", "What it is", "Why they would choose it"],
      [
          ["A — The engines",
           "Keep Jasiri. Licence the Internal Report API and the credit scoring engines. Statement in, "
           "score + starting limit + matched product + reason codes out.",
           "Lowest risk, fastest to value, no migration. The brain without the surgery."],
          ["B — The front door",
           "Keep Jasiri as the book of record. We build identity-first onboarding and the customer "
           "portal, feeding into it.",
           "Fixes turnaround — the actual problem — without touching what already works."],
          ["C — The platform",
           "Move onto BirgenAI LMS. Mular-branded portal, the Hub pipeline, the full engine suite, "
           "and the closed learning loop from day one.",
           "The only option where their own model eventually decides their own book."],
      ],
      widths=[3.0, 6.6, 6.2])

h2(doc, "The close")
say(doc, "I am not asking you to choose today. I am asking for one thing: let me run the scoring "
         "engine in shadow against your last ninety days of decisions. It changes nothing, it "
         "risks nothing, and it will tell you more about whether this is worth doing than anything "
         "I can say in this room.")

note(doc, "This is a near-unrefusable ask. It is free, reversible, invisible to their customers, and "
          "it puts your engine inside their data. Whoever says yes to this has effectively started.")

h2(doc, "If they say yes")
say(doc, "Then I need three things: a read-only extract of the last ninety days, one technical "
         "contact, and a date two weeks out to look at the results together. I will send that in "
         "writing this afternoon.")

note(doc, "Name the date before you leave the room. A meeting with no next date is a meeting that ended.")

# ── 9 ────────────────────────────────────────────────────────────────────────
h1(doc, "9", "Language — what to reach for, what to avoid")

table(doc,
      ["Do not say", "Say instead", "Why"],
      [
          ["“Your system is broken / bad / outdated.”",
           "“Jasiri is capable and is being held back by three things.”",
           "Somebody in the room chose it. Attack the choice and you lose the chooser."],
          ["“AI-powered.”",
           "“It reads their statement and gives you a limit you can explain.”",
           "Everyone says AI. Nobody says explain. The second one sells."],
          ["“We can do anything.”",
           "“Here is what it does, and here is where it stops.”",
           "Stating a limit is the fastest way to be believed about everything else."],
          ["“It is fully automated.”",
           "“It moves the easy decisions off your officers' desks.”",
           "Lenders hear 'automated' as 'loss of control'. Never trigger that."],
          ["“Trust me.”",
           "“Run it in shadow and check it yourself.”",
           "Never ask for trust you can replace with evidence."],
          ["“I hacked your system.”",
           "“During normal testing I was able to… and I disclosed it privately.”",
           "Same fact, opposite feeling. One makes you a threat, the other a professional."],
          ["“Cheap / affordable.”",
           "“It pays for itself in turnaround.”",
           "Price framing invites price negotiation. Value framing invites scope."],
          ["“Our model is 82% accurate.”",
           "“Here is the evidence we have so far, and here is what it is not yet enough to prove.”",
           "A precise number invites a fight. Honest limits end one."],
      ],
      widths=[4.6, 5.6, 5.6])

h2(doc, "Three phrases worth memorising")
bullet(doc, "“You are paying full price for decisions that make themselves.”", size=11)
bullet(doc, "“Same person, two screens, one status.”", size=11)
bullet(doc, "“A number you cannot check is a number you cannot act on.”", size=11)

# ── 10 ───────────────────────────────────────────────────────────────────────
h1(doc, "10", "One-page crib — print this bit")

para(doc, "Open", 11, bold=True, color=GREEN, before=4, after=2)
bullet(doc, "Thank them for access. Praise the analytics honestly.")
bullet(doc, "Three things holding it back: onboarding friction, unfinished areas, security exposures. None fatal.")

para(doc, "Headline", 11, bold=True, color=GREEN, before=8, after=2)
bullet(doc, "Turnaround is the money. IPRS identity + statement scoring with reason codes.")
bullet(doc, "“You are paying full price for decisions that make themselves.” Then pause.")

para(doc, "Security", 11, bold=True, color=GREEN, before=8, after=2)
bullet(doc, "Say it once, calmly. Took nothing, changed nothing, disclosed privately.")
bullet(doc, "“It is not a reason to hire me. It is just the right thing to do.”")

para(doc, "Demo", 11, bold=True, color=GREEN, before=8, after=2)
bullet(doc, "Hub → Mular → onboard → crunch → both screens → sign on phone → PIN reset → AI → report → loop.")
bullet(doc, "Stop talking at the two-screens moment and at the signing moment.")

para(doc, "Science", 11, bold=True, color=GREEN, before=8, after=2)
bullet(doc, "300 outcomes, then their own model decides their own book.")
bullet(doc, "“Below that, nobody can honestly claim otherwise. Not me, not anyone.”")

para(doc, "Ask", 11, bold=True, color=GREEN, before=8, after=2)
bullet(doc, "Three doors: engines / front door / platform. Smallest first.")
bullet(doc, "Close on the shadow run. Ninety days of data, one contact, a date two weeks out.")
bullet(doc, "Do not leave without the date.", color=RED)

doc.add_paragraph()
rule(doc, "DDDDDD", 6)
para(doc, "Emmanuel Birgen  ·  IT Consultant, Serve Well Co.  ·  Nairobi", 9, color=MUTED, after=1)
para(doc, "Companion to the Jasiri LMS Independent Assessment, 24 July 2026. Confidential.",
     8.5, color=FAINT, italic=True)

os.makedirs("reports", exist_ok=True)
doc.save(OUT)
print("wrote " + OUT)
