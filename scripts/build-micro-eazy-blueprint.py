# -*- coding: utf-8 -*-
"""
Build the Micro Eazy Ecosystem Blueprint (.docx).

Two audiences in one document, and the split is deliberate: Sections 1–4 are what
a Micromart board member should be able to read cold, and Sections 5–12 are the
engineering contract we build against. Emerald and amber throughout — the Micro
Eazy palette — so it sits in its own family rather than borrowing Mular's navy.

Source of truth is docs/MICRO-EAZY-ECOSYSTEM.md; this renders the shareable copy.

    python scripts/build-micro-eazy-blueprint.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEEP = RGBColor(0x06, 0x5F, 0x46)   # emerald, deep — headings
GREEN = RGBColor(0x10, 0xB9, 0x81)  # emerald, bright — accents
AMBER = RGBColor(0xB4, 0x7C, 0x00)  # amber, print-safe — emphasis
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
FAINT = RGBColor(0x8A, 0x8A, 0x8A)
RED = RGBColor(0xB0, 0x2A, 0x2A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

DEEP_HEX = "065F46"
OUT = os.path.join("reports", "Micro-Eazy-Ecosystem-Blueprint.docx")


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def rule(doc, color=DEEP_HEX, size=12):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    b.append(bottom)
    pPr.append(b)
    p.paragraph_format.space_after = Pt(6)
    return p


def para(doc, text="", size=10.5, bold=False, italic=False, color=INK,
         before=0, after=4, indent=0, font="Calibri", align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent:
        pf.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.name = font
    return p


def quote(doc, text, indent=0.6, color="10B981"):
    """A line meant to be said out loud, or a claim we are standing behind."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(indent)
    pf.space_before = Pt(4)
    pf.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    b.append(left)
    pPr.append(b)
    r = p.add_run(u"“" + text + u"”")
    r.font.size = Pt(11)
    r.font.name = "Georgia"
    r.font.color.rgb = INK
    return p


def note(doc, text, indent=0.6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(indent)
    pf.space_before = Pt(0)
    pf.space_after = Pt(7)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = FAINT
    r.font.name = "Calibri"
    return p


def bullet(doc, text, size=10.5, indent=0.85, color=INK, bold_lead=None, marker=u"•   "):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(indent)
    pf.space_after = Pt(3)
    r0 = p.add_run(marker)
    r0.font.size = Pt(size)
    r0.font.color.rgb = GREEN
    r0.font.name = "Calibri"
    if bold_lead:
        rb = p.add_run(bold_lead)
        rb.font.size = Pt(size)
        rb.font.bold = True
        rb.font.color.rgb = INK
        rb.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def mono(doc, lines, size=8.5, indent=0.6, fill="F4F7F5"):
    """A fixed-width block: schema, a flow diagram, a worked number."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    shade(cell, fill)
    cell.text = ""
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(0.2)
        r = p.add_run(line)
        r.font.size = Pt(size)
        r.font.name = "Consolas"
        r.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def h1(doc, num, text, page_break=True):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(num)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = GREEN
    r.font.name = "Calibri"
    r2 = p.add_run("   " + text)
    r2.font.size = Pt(17)
    r2.font.bold = True
    r2.font.color.rgb = DEEP
    r2.font.name = "Calibri"
    rule(doc, "DDDDDD", 6)


def h2(doc, text, color=DEEP):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def table(doc, headers, rows, widths=None, head_fill=DEEP_HEX, bold_first=True, size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        shade(hdr[i], head_fill)
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(htxt)
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(size)
            r.font.color.rgb = INK
            r.font.name = "Calibri"
            if i == 0 and bold_first:
                r.font.bold = True
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ═════════════════════════════════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)

# ── Cover ────────────────────────────────────────────────────────────────────
para(doc, "BIRGENAI  ·  PRODUCT BLUEPRINT", 10, bold=True, color=GREEN, after=2)
para(doc, "Micro Eazy", 38, bold=True, color=DEEP, after=0)
para(doc, "A two-sided credit ecosystem — and the plan to put "
          "Micromart at the centre of it", 13, color=MUTED, after=14)
rule(doc)
para(doc, "Emmanuel Birgen  ·  Founder, BirgenAI", 10.5, bold=True, after=1)
para(doc, "Anchor lender: Micromart Africa Limited  ·  Board presentation: week of 10 August 2026",
     9.5, color=MUTED, after=1)
para(doc, "5 August 2026  ·  Version 1.0  ·  Grounded in a live read of the platform database",
     9.5, color=MUTED, after=16)

para(doc, "HOW TO READ THIS DOCUMENT", 9, bold=True, color=GREEN, after=3)
bullet(doc, "read cold by anyone on the Micromart board. No engineering knowledge assumed.",
       bold_lead="Sections 1–4 can be ")
bullet(doc, "the engineering contract — what gets built, in what order, and against which files.",
       bold_lead="Sections 5–12 are ")
bullet(doc, "are things we are standing behind, or words to say in the room.",
       bold_lead="Green-barred lines ")
bullet(doc, "are the reasoning — background, never to be read aloud.", bold_lead="Grey italics ")
bullet(doc, "Every claim about what already exists was verified against the running system on "
            "5 August 2026, not recalled. Section 2 shows the evidence.")

doc.add_paragraph()
para(doc, "THE ONE THING TO GET RIGHT", 9, bold=True, color=RED, after=3)
para(doc, "Micromart has already built Micro Eazy. It is live, correctly configured, and it has "
          "two loans on it. We are not selling them a product, a system, or a migration. We are "
          "selling them the pipeline for the product they already own — and asking for nothing "
          "back except the shelf. Every technical decision in this document exists to make that "
          "sentence true and safe. If a thread stops serving it, drop the thread.", 10.5, after=6)

# ── 1 ────────────────────────────────────────────────────────────────────────
h1(doc, "1", "The proposition")

h2(doc, "What Micro Eazy is")
para(doc, "Micro Eazy is a two-sided credit ecosystem.")
bullet(doc, "licensed lenders run their entire business on the BirgenAI Super LMS — origination, "
            "scoring, workflow, disbursement, collections, field operations, accounting, HR and the "
            "call centre — under one sign-on.", bold_lead="On the lender side, ")
bullet(doc, "Kenyans install one app, Micro Eazy, from the BirgenAI Hub app store, verify themselves "
            "once, and are routed to a licensed lender who funds them.", bold_lead="On the customer side, ")
para(doc, "BirgenAI never lends. The licensed lender is always lender-of-record, named on every "
          "money screen, every offer and every agreement. What BirgenAI owns is the customer "
          "relationship, the verified identity, the intelligence and the rails.", before=4)

h2(doc, "Why Micromart, and why now")
para(doc, "Micromart’s own product screen makes the argument better than any slide could. "
          "Micro Eazy Monthly is live and correctly configured — 22% flat per month, two months, "
          "6% processing fee, minimum credit score 500 — and it carries "
          "2 loans and KES 144,000 outstanding.")
quote(doc, "You have built Micro Eazy. We have built the machine that fills it. Same name, same "
           "product, same workflow you already approved — and from next week, a stream of "
           "verified, scored, geo-pinned, CRB-checked customers arriving in it every day.")
note(doc, "Getting a licensed lender to pilot a credit platform is the hardest door in Kenyan "
          "fintech. It is open. Everything below is about walking through it without breaking a "
          "single thing they already run.")

h2(doc, "What Micromart is being asked for")
bullet(doc, "Keep Micro Eazy and Micro Eazy Monthly active on their shelf.")
bullet(doc, "Accept applications that arrive in their existing Micro Eazy approval workflow, "
            "reviewed by their own officers, in their own system.")
bullet(doc, "Agree a launch appetite — loans per day and maximum exposure.")
para(doc, "That is the whole ask. No migration, no new system for their staff to learn, no change "
          "to how money leaves their account.", before=4, bold=True)

# ── 2 ────────────────────────────────────────────────────────────────────────
h1(doc, "2", "Where we actually are")

note(doc, "This section is evidence, not a status claim. Every line was probed against the live "
          "platform database and the running codebase on 5 August 2026.")

h2(doc, "2.1  The lender realm — Super LMS")
para(doc, "Next.js 16 · Prisma 7 · PostgreSQL with row-level security enforced inside the "
          "database, so tenant isolation is a database guarantee rather than a code-review promise. "
          "62 data models. Thirteen console sections, roughly forty-five screens.")
table(doc,
      ["Capability", "State"],
      [
          ["Multi-tenant schema, RLS isolation", "BUILT"],
          ["Decision engine — 7 stages, reason codes on every stage, reproducible", "BUILT"],
          ["Per-lender credit policy, declarative and versioned", "BUILT"],
          ["Product builder, versioning, published eligibility rules", "BUILT"],
          ["Workflow engine — stage tree, access tiers, OTP approvals, finalize caps", "BUILT"],
          ["M-Pesa statement cruncher and Internal Report", "BUILT"],
          ["Scoring: v2 bespoke (AUC 0.822), v3.1.1 pooled (0.823), v1 behavioural monitor", "LIVE"],
          ["Metropol CRB integration", "BUILT · test credentials"],
          ["ServiceSuite bridge — posts into the lender’s own live workflow", "LIVE"],
          ["Collections, promises-to-pay, call logs, tickets", "BUILT"],
          ["Field ops — geo pins, route planner, disbursement location gate", "BUILT"],
          ["Riri assistant — role, book and customer aware", "BUILT"],
          ["Connected Suite SSO — Lending, Portal, HR, Accounting, Call-Centre", "BUILT"],
          ["Billing, usage metering, invoices, entitlements", "BUILT"],
      ],
      widths=[12.6, 4.2])

h2(doc, "2.2  Micromart, as the platform sees it today")
mono(doc, [
    "micromart    BRIDGED   ACTIVE   PREMIUM   Micromart Africa",
    "",
    "  borrowers    162        loans       199      applications  50",
    "  offers        15        staff        17      branches       9",
    "  products       5  (1 active)         workflows              2",
    "  integrations  MPESA_STK = CONFIGURED   ·   CRB = CONFIGURED",
    "  bridge        live — posts via sp_InsertLoan, channel 7",
])
para(doc, "Micro Eazy and Micro Eazy Monthly do not yet exist on our side. That is a configuration "
          "task, not an engineering one — which is precisely what the product builder is for.",
     before=2)

h2(doc, "2.3  The customer realm — what exists, and what is missing")
para(doc, "The borrower surfaces are real and working. They are simply not yet an app.")
table(doc,
      ["Customer surface", "State"],
      [
          ["Branded application wizard — phone, consent, crunch, score, offer", "BUILT"],
          ["OTP door, national-ID + PIN door, session", "BUILT"],
          ["Crunch theatre — the statement analysis, watchable", "BUILT"],
          ["Offer, full schedule, e-signature", "BUILT"],
          ["My loan, balance, Pay Now by STK push", "BUILT"],
          ["Auto-repay (M-Pesa Ratiba)", "BUILT"],
          ["Internal Report sold to the customer", "BUILT"],
          ["Guarantor invitation and consent", "BUILT"],
          ["Installable app — manifest, icons, service worker", "MISSING"],
          ["Push notifications, notification inbox", "MISSING"],
          ["Offline shell and background sync", "MISSING"],
          ["“Why was I declined” and “how do I fix it”", "MISSING"],
          ["Limit ladder and graduation progress", "MISSING"],
          ["Rewards, tiers, early-settlement rebate", "MISSING"],
          ["Multi-lender routing — the Exchange", "MISSING — the new build"],
      ],
      widths=[12.6, 4.2])

h2(doc, "2.4  How Micro Eazy appears on the BirgenAI Hub")
para(doc, "The Hub app store reads a single table. Making Micro Eazy appear on the "
          "“Loans & Credit” shelf is three steps, and the first one is the logo:")
bullet(doc, "the logo at public/apps/icons/micro-eazy.png, 512×512, transparent.",
       marker=u"1.  ", bold_lead="Drop ")
bullet(doc, "in the apps seed — category LENDING, featured, sort order 0, "
            "URL microeazy.birgenai.com.", marker=u"2.  ", bold_lead="Add one row ")
bullet(doc, "the seed. Micro Eazy is now the first tile on the shelf.",
       marker=u"3.  ", bold_lead="Re-run ")

# ── 3 ────────────────────────────────────────────────────────────────────────
h1(doc, "3", "The three decisions, locked")

table(doc,
      ["Decision", "What we chose", "Why"],
      [
          ["D1  ·  The app",
           "ONE customer PWA at microeazy.birgenai.com. After routing, its chrome repaints to the "
           "assigned lender. Lender subdomains survive as branded doors into the same installed app.",
           "One manifest, one home-screen icon, one push channel, one install base — owned by "
           "BirgenAI. Lender #2 costs zero customer re-installs. Per-lender apps fragment the "
           "install base and hand it to the lender."],
          ["D2  ·  The brand",
           "Co-branded. “Micro Eazy · Funded and serviced by Micromart Africa Ltd · "
           "Powered by BirgenAI” on every money screen, offer, agreement and SMS.",
           "It is the honest regulatory position — BirgenAI never lends. It names Micromart "
           "where it legally matters. And it keeps the consumer brand portable to lender #2."],
          ["D3  ·  The product",
           "Micro Eazy mirrors Micromart’s terms exactly — same names, same rates, same "
           "fee, same workflow names.",
           "The board must recognise their own product on screen. Zero configuration argument in "
           "the room. Divergence is a later, negotiated conversation."],
      ],
      widths=[3.4, 6.4, 7.0], size=9)

# ── 4 ────────────────────────────────────────────────────────────────────────
h1(doc, "4", "The Trust Contract")

para(doc, "This is what makes Micro Eazy an ecosystem rather than a lead-generation funnel. Both "
          "sides owe each other something, and both sets of obligations are enforced in code, not "
          "promised in a brochure. This is the section to read aloud to the board.")

h2(doc, "4.1  What the customer gives")
para(doc, "Stated plainly, once, before anything is collected. Mandatory to receive money through "
          "Micro Eazy:")
bullet(doc, "verified identity — national ID capture, selfie with liveness, IPRS cross-check")
bullet(doc, "consent to CRB check, M-Pesa statement analysis and automated scoring — versioned "
            "and IP-stamped")
bullet(doc, "a location pin that can actually be visited — already a hard disbursement gate")
bullet(doc, "device fingerprint, as a fraud signal")
bullet(doc, "repayment on the agreed schedule")

h2(doc, "4.2  What the ecosystem owes back — seven enforceable promises")
table(doc,
      ["Promise", "How it is enforced"],
      [
          ["No silent decline",
           "Engine reason codes surfaced verbatim to the customer, in English and Kiswahili, with a "
           "named path to fix each one."],
          ["A visible ladder",
           "“Rung 2 of 7. Clear this on time → KES 15,000. Clear it early → KES 18,000.”"],
          ["Pay early, pay less — priced live",
           "Early-settlement rebate as a slider on the loan card, not a phone call to the branch."],
          ["Your data, your report",
           "The Internal Report — the same analysis the lender buys — free to the customer "
           "once per cycle."],
          ["Portability",
           "Good behaviour travels to the next lender in the ecosystem under a stated legal basis. "
           "The reward is not locked inside one lender’s book."],
          ["A human answers appeals",
           "Adverse decisions are never fully automatic. The appeal clock is displayed and counted down."],
          ["Purpose and retention, always stated",
           "Retention windows per data class; export and deletion workflows per organisation."],
      ],
      widths=[5.2, 11.6], size=9.5)

h2(doc, "4.3  Rewards — the ladder up")
para(doc, "Clean cycle → limit graduation · rate step-down · processing-fee waiver at "
          "rung four · instant re-borrow that skips the officer stages the lender marked "
          "skippable · early-settlement rebate · Micro Eazy Gold, which routes to the "
          "best-priced listing first · referral credit.")

h2(doc, "4.4  Consequences — proportionate, disclosed up front, never a surprise")
para(doc, "Missed instalment → reminder ladder at T-3, T-1 and T0 → promise-to-pay → "
          "agent call → geo-pinned field visit → limit freeze → limit reduction → "
          "rate step-up → CRB listing at the disclosed threshold, by the lender-of-record "
          "→ ecosystem stop-flag.")
quote(doc, "Every rung is written into the agreement the customer signs, and every rung is visible "
           "in the app before it is reached. A consequence the customer was warned about is "
           "collections. A consequence they were not is a complaint to the regulator.")

h2(doc, "4.5  Responsible AI — non-negotiable, ecosystem-wide")
bullet(doc, "Every model output carries its model version, input hash and reason codes, persisted.")
bullet(doc, "Server-side recompute only. The client is never trusted with a score.")
bullet(doc, "No fully automatic adverse decision beyond the lender’s own disclosed floor.")
bullet(doc, "Every decision is reproducible — “why was I declined in March?” is "
            "answerable in March’s terms, because the policy version, product version, inputs "
            "and full stage trace are all stored.")
bullet(doc, "Cross-border minimisation — features and aggregates leave, raw personal data does not.")
bullet(doc, "Bias monitoring on the closed ML loop: the 300-outcome gate, Wilson intervals, "
            "priced errors.")

# ── 5 ────────────────────────────────────────────────────────────────────────
h1(doc, "5", "Architecture — two realms, one spine")

mono(doc, [
    "                    BIRGENAI HUB  (birgenai.com)",
    "             App store  >  Loans & Credit  >  [ Micro Eazy ]",
    "                              |  install / launch (SSO)",
    "                              v",
    "  +---------------------+  +--------------------+  +----------------------------+",
    "  |  REALM B  CUSTOMER  |  |     THE SPINE      |  |  REALM A  LENDER           |",
    "  | microeazy.birgenai  |  | MICRO EAZY EXCHANGE|  |  Super LMS console         |",
    "  | installable PWA     |  |                    |  |                            |",
    "  |                     |  |  . Listings        |  |  . products & workflows    |",
    "  | . one identity      |<>|  . Allocation      |<>|  . officer queues          |",
    "  | . one KYC, reusable |  |    policy          |  |  . disbursement / float    |",
    "  | . apply/offer/sign  |  |  . Appetite ledger |  |  . collections / field     |",
    "  | . pay / auto-repay  |  |  . RoFR + SLA      |  |  . Riri / analytics        |",
    "  | . limit ladder      |  |  . Routing record  |  |  . HR / Accounts / Calls   |",
    "  | . why-declined      |  |    (reason codes)  |  |                            |",
    "  | . Internal Report   |  |                    |  |  ServiceSuite BRIDGE ------+--> Micromart",
    "  | . rewards / tier    |  | -- TRUST CONTRACT -|  |  sp_InsertLoan -> workflow |    live book",
    "  +---------------------+  +--------------------+  +----------------------------+",
    "",
    "        SHARED INTELLIGENCE  (API-first, separable, embeddable)",
    "   Statement Cruncher . Internal Score . CRB . Decision Engine .",
    "   Behavioural Monitor . Riri",
], size=7.6)

para(doc, "Realm A is roughly 80% built. The spine, and Realm B’s shell, are the new work.",
     before=4, bold=True)

# ── 6 ────────────────────────────────────────────────────────────────────────
h1(doc, "6", "The Micro Eazy Exchange — the actual invention")

para(doc, "Everything else in this document is assembly. This is the new product.")
para(doc, "Today the decision engine’s routing stage answers “who inside this lender may "
          "approve?” It does not answer “which lender gets this customer?” — "
          "because until now there has only ever been one.")

h2(doc, "6.1  Design principle: a pre-pass, not a rewrite")
para(doc, "The decision engine is pure, reproducible and covered by parity tests. It does not "
          "change. The Exchange runs before it and picks the lender; the existing per-lender engine "
          "then runs unchanged inside that lender’s policy.")
mono(doc, [
    "  applicant",
    "     |",
    "     v",
    "  +--------------------------------------------+",
    "  |  EXCHANGE PRE-PASS                         |",
    "  |  1. eligible listings for this applicant   |",
    "  |  2. appetite / quota check per lender      |",
    "  |  3. allocation policy applied              |",
    "  |  4. right-of-first-refusal + SLA clock     |",
    "  |  -> { orgId, listingId, reasons[] }        |",
    "  +---------------------+----------------------+",
    "                        v",
    "     candidatesFor(orgId) -> engine.decide(...)   <-- UNCHANGED",
    "                        v",
    "     offer . workflow . disbursement (that lender's rails)",
])

h2(doc, "6.2  The five allocation modes")
table(doc,
      ["Mode", "What it does", "When it turns on"],
      [
          ["SOLE", "One lender takes everything.", "Launch — Micromart."],
          ["WEIGHTED", "Round-robin by capital-share weights.",
           "Lender #2 and #3 arrive with different balance sheets."],
          ["CAPACITY_FIRST", "Route to whoever has remaining appetite today.",
           "Protects lenders from flooding, customers from silent queues."],
          ["BEST_FIT", "Price the applicant against every live listing; the customer gets the best "
                       "affordable offer and sees why.",
           "The endgame — this is what makes Micro Eazy a market rather than a funnel."],
          ["WATERFALL", "First-look lender holds right of first refusal for N minutes, then the "
                        "customer cascades.",
           "Sells premium first-look as a paid tier."],
      ],
      widths=[3.0, 7.4, 6.4], size=9)
quote(doc, "SOLE at launch. The other four ship in the code from day one, switched off — which "
           "is exactly what lets us tell the Micromart board “you are our only lender” and "
           "mean it, without rebuilding anything when that changes.")

h2(doc, "6.3  Every routing explains itself")
para(doc, "The routing record reuses the engine’s reason-code shape, so the customer-facing "
          "“why” and the regulator-facing audit trail are the same object:")
mono(doc, [
    "LENDER_AWARDED     Micromart Africa     Sole lender for Micro Eazy in this period",
    "LISTING_MATCH      Micro Eazy Monthly   KES 5,000-100,000 fits your assessed limit",
    "                                        of KES 25,000",
    "LISTING_EXCLUDED   Micro Eazy (weekly)  Your income cycle is monthly; weekly terms",
    "                                        were not offered",
    "APPETITE_OK        Micromart            Within today's remaining capacity",
])
note(doc, "That last block is also a sales asset. It is the report you put in front of a lender to "
          "sell them a bigger quota.")

# ── 7 ────────────────────────────────────────────────────────────────────────
h1(doc, "7", "Realm B — the Micro Eazy app")

para(doc, "Same codebase, new route group and host. No new repository.")

h2(doc, "7.1  Ten of sixteen screens already exist")
table(doc,
      ["Screen", "Work"],
      [
          ["Splash and install coaching — Android prompt, iOS add-to-home-screen", "NEW"],
          ["Phone + OTP door", "re-shell"],
          ["National ID + PIN door", "re-shell"],
          ["Consent — granular, versioned", "re-shell"],
          ["KYC — ID capture, quality gates, selfie, liveness, face match", "finish"],
          ["Crunch theatre", "re-shell"],
          ["Offer, full schedule, e-signature", "re-shell"],
          ["Home — current loan, due today, pay slider", "re-shell"],
          ["Auto-repay (Ratiba)", "re-shell"],
          ["Internal Report", "re-shell"],
          ["Why declined, and how to fix it", "NEW"],
          ["Limit ladder and graduation", "NEW"],
          ["Rewards and tier", "NEW"],
          ["Notification inbox and push", "NEW"],
          ["Offline shell and background sync", "NEW"],
          ["Support — Riri chat and tickets", "port"],
      ],
      widths=[13.0, 3.8])
para(doc, "This is why a two-week path to a working demo is realistic rather than delusional: the "
          "app is mostly a re-shell of code that already works, plus six new screens.", before=2)

h2(doc, "7.2  The install path, end to end")
mono(doc, [
    "  Hub app store > Loans & Credit > [Micro Eazy]  -- install -->  microeazy.birgenai.com",
    "        |                                                              |",
    "        +-- SSO: BirgenAI ID already signed in ------------------------+",
    "                              |",
    "                    Android: install prompt -> home-screen icon",
    "                    iOS: Safari share-sheet coaching card",
    "                              |",
    "               standalone app, Micro Eazy icon, no browser chrome",
])

h2(doc, "7.3  Specification")
bullet(doc, "standalone display, icons at 192, 512 and 512-maskable, theme colour, portrait "
            "orientation, and shortcuts for Pay Now, My Loan and Apply.", bold_lead="Manifest: ")
bullet(doc, "app-shell precache, network-first for the API, and background sync for pending "
            "repayments and consent submissions — Kenyan network reality, not a nicety.",
       bold_lead="Service worker: ")
bullet(doc, "approved, disbursed, payment received, due at T-3 / T-1 / T0, limit increased.",
       bold_lead="Push: ")
bullet(doc, "mobile-first Android. Every touch target at least 44 pixels. The whole flow works "
            "one-handed on a 360-pixel viewport.", bold_lead="Layout: ")

# ── 8 ────────────────────────────────────────────────────────────────────────
h1(doc, "8", "Micromart — the exact configuration")

h2(doc, "8.1  Micro Eazy Monthly, mirrored from their live screen")
table(doc,
      ["Field", "Value"],
      [
          ["Principal", "KES 5,000 – 100,000"],
          ["Interest method", "Flat rate"],
          ["Interest rate", "22.00% per month"],
          ["Repayment", "2 (Month)"],
          ["Rollover penalty", "20.00%"],
          ["New loans", "Approval required · workflow “Micro Eazy”"],
          ["Repeat loans", "Approval required · workflow “Micro Eazy”"],
          ["Guarantor", "Not required · in-active guarantors cannot borrow"],
          ["Security", "Not required"],
          ["Minimum credit score", "500.00"],
          ["Minimum loan limit", "KES 5,000"],
          ["Charge", "PROCESSING FEE (PF) · before disbursement · 6.00% · capped "
                     "KES 650–6,000 · range KES 5,000–100,000 · mandatory · active"],
      ],
      widths=[5.0, 11.8])

h2(doc, "8.2  The worked example the board will do in their heads")
note(doc, "Get this right on the first screen they see. If the arithmetic on our offer card does "
          "not match the arithmetic in their heads, nothing else in the demo lands.")
mono(doc, [
    "  Principal                                    KES  25,000",
    "  Processing fee  6%  (inside the 650-6,000 cap)    1,500   deducted before disbursement",
    "  ------------------------------------------------------",
    "  NET TO CUSTOMER                              KES  23,500",
    "",
    "  Interest  22% flat per month x 2 months on 25,000  11,000",
    "  ------------------------------------------------------",
    "  TOTAL REPAYABLE                              KES  36,000",
    "  Monthly instalment                           KES  18,000  x 2",
    "",
    "  Rollover penalty                             20%  if rolled",
])

h2(doc, "8.3  Micro Eazy — the base product — OPEN")
para(doc, "Micromart has two Micro Eazy products; only Monthly’s screen was supplied. Needed "
          "before the demo: principal range, interest rate and unit, repayment count and unit, "
          "rollover penalty, minimum credit score, minimum loan limit, charge structure and "
          "workflow name. Until it arrives, Micro Eazy Monthly carries the demo alone.")

h2(doc, "8.4  What gets configured, not coded")
bullet(doc, "Two product records in the Micromart organisation, exact mirrors, set active — "
            "and for a bridged lender, activating a product is what puts it on their live shelf.")
bullet(doc, "One charge record: processing fee, 6%, minimum 650, maximum 6,000, before "
            "disbursement, mandatory.")
bullet(doc, "One workflow titled “Micro Eazy”, stages mirroring theirs, mapped to their "
            "ServiceSuite approval workflow.")
bullet(doc, "A credit policy whose automatic-decline floor matches their minimum score of 500, and "
            "whose automatic-approval band is set conservatively so that every launch loan is "
            "officer-reviewed. Human in the loop, and the board sees their officers still in control.")
bullet(doc, "Two marketplace listings, marked exclusive.")
bullet(doc, "One allocation policy: SOLE, awarding to Micromart.")

# ── 9 ────────────────────────────────────────────────────────────────────────
h1(doc, "9", "The demo — what happens in the room")

quote(doc, "Do not present slides. Present a loan.")

steps = [
    ("1", "Open birgenai.com on the projector. App store, Loans & Credit. Micro Eazy is the first "
          "tile, carrying the logo they are about to see everywhere."),
    ("2", "Install it on a real Android phone, on stage. The icon lands on the home screen."),
    ("3", "Apply as a real customer: phone, OTP, consent, ID and selfie, M-Pesa statement. Let them "
          "watch the crunch theatre run. CRB pull. Decision."),
    ("4", "The offer appears — “Micro Eazy Monthly · KES 25,000 · 2 months · "
          "KES 18,000 per month · Funded and serviced by Micromart Africa Ltd” — and "
          "beside it, the reason codes that produced it."),
    ("5", "Accept. E-sign by OTP."),
    ("6", "Turn the projector to MICROMART’S OWN SERVICESUITE SCREEN. The loan is sitting in "
          "their Micro Eazy workflow, at Officer Review, unapproved. Their loan. Their workflow. "
          "Their officer."),
    ("7", "A Micromart officer approves it live, in their own system."),
    ("8", "Back to the phone: disbursed. SMS receipt. Balance showing."),
    ("9", "Repay by STK push from the phone. The balance drops. The limit ladder moves — "
          "“Rung 2 → KES 30,000 next.”"),
    ("10", "Show the customer’s “why” screen and their Internal Report. Then show the "
           "lender’s side: the pipeline, the attribution, the reason trace."),
]
for num, text in steps:
    bullet(doc, text, marker=num.rjust(2) + ".  ")

doc.add_paragraph()
para(doc, "Then the sentence that closes it:", bold=True, after=2)
quote(doc, "Your Micro Eazy Monthly has two loans on it. Everything you just watched took four "
           "minutes and did not touch a single line of your system. Give us the shelf, and we fill it.")

# ── 10 ───────────────────────────────────────────────────────────────────────
h1(doc, "10", "Build plan")

h2(doc, "Sprint 0 — demo-critical (this week → the board meeting)")
table(doc,
      ["#", "Task"],
      [
          ["0.1", "Micro Eazy logo — three concepts, SVG and PNG at every needed size"],
          ["0.2", "Seed Micro Eazy and Micro Eazy Monthly into the Micromart organisation, exact mirror"],
          ["0.3", "Mirror the “Micro Eazy” workflow; map it to their ServiceSuite approval workflow"],
          ["0.4", "Exchange v1 — schema, allocation module, SOLE policy awarding to Micromart"],
          ["0.5", "PWA — manifest, icons, service worker, install prompt, iOS coaching card"],
          ["0.6", "microeazy.birgenai.com host routing and co-branded chrome"],
          ["0.7", "Hub tile — icon and seed row, LENDING category, featured, first position"],
          ["0.8", "Customer “why declined / how to fix” screen, from the reason codes that already exist"],
          ["0.9", "Limit-ladder screen, from the graduation events that already exist"],
          ["0.10", "End-to-end rehearsal on the real Micromart bridge — twice, with a reversible test loan"],
          ["0.11", "Fix the product-name typo “MIROMART FINTECH”"],
      ],
      widths=[1.6, 15.2])

h2(doc, "Sprint 1 — the Trust Contract made real (two weeks after the demo)")
para(doc, "Push notifications · notification inbox · offline shell and background sync · "
          "rewards and tiers · early-settlement rebate slider · appeal flow with a visible "
          "SLA clock · Internal Report free once per cycle · full Kiswahili pass on every "
          "new screen.")

h2(doc, "Sprint 2 — the Exchange opens (weeks three to six)")
para(doc, "Lender marketplace console · appetite ledger and quota enforcement · WEIGHTED "
          "and CAPACITY_FIRST modes · right-of-first-refusal with SLA lapse · attribution "
          "reporting and lender invoicing · lender self-onboarding to the shelf.")

h2(doc, "Sprint 3 — one workspace (weeks six to nine)")
para(doc, "The Connected Suite as the lender’s whole back office — HR, Accounting, "
          "Call-Centre, Analytics Studio — under one sign-on, with the app drawer intact. "
          "Riri across all five.")

h2(doc, "Sprint 4 — scale (month three onward)")
para(doc, "BEST_FIT and WATERFALL modes · per-lender model calibration · lender #2 and #3 "
          "· national coverage via the risk map · direct IPRS · WhatsApp channel.")

# ── 11 ───────────────────────────────────────────────────────────────────────
h1(doc, "11", "Risks, and what we do about them")

table(doc,
      ["Risk", "The reality", "What we do"],
      [
          ["Live posting is armed",
           "Posting runs against Micromart’s production system today.",
           "Every rehearsal uses a reversible test borrower and a documented rollback. Never "
           "rehearse blind."],
          ["Real SMS credentials sit in the environment file",
           "A test run can send real messages and spend real money.",
           "Blank them in every verification and rehearsal script. Standing rule."],
          ["CRB production keys pending",
           "Only test credentials today.",
           "Demo on test credentials — the integration is already proven green. Chase the "
           "production keys this week."],
          ["The bridge is environment-driven, not vaulted",
           "There is no per-organisation integration record for the ServiceSuite bridge.",
           "Move the bridge configuration into the encrypted vault so it is per-lender and "
           "console-managed. Sprint 1."],
          ["Base product spec missing",
           "Only Micro Eazy Monthly’s screen was supplied.",
           "Ask Micromart for the second screen. Monthly carries the demo alone if it does not arrive."],
          ["Sole-lender concentration",
           "One lender’s appetite caps the whole ecosystem.",
           "The other four allocation modes ship from day one, switched off. Lender #2 is a "
           "configuration change, not a project."],
          ["Regulatory posture",
           "BirgenAI is not a licensed lender.",
           "Lender-of-record named on every money screen. Data-protection registration and impact "
           "assessment before external scale."],
      ],
      widths=[3.4, 6.2, 7.2], size=9)

# ── 12 ───────────────────────────────────────────────────────────────────────
h1(doc, "12", "Open items — founder")

items = [
    ("Micro Eazy base product spec", "The second product screen from Micromart."),
    ("The “Micro Eazy” workflow stages", "Names, order, and who approves at each."),
    ("microeazy.birgenai.com", "DNS record and Vercel domain."),
    ("Web Push keypair", "Generated once, stored in the vault."),
    ("CRB production keys", "Chase this week."),
    ("Micromart’s launch appetite",
     "Loans per day and maximum exposure. A commercial number — it belongs in the board "
     "conversation, not in a config file we choose."),
    ("Commercial model",
     "Per-loan origination fee, revenue share, or subscription plus usage. Needed before the "
     "attribution reporting in Sprint 2 has anything to invoice against."),
]
table(doc, ["Item", "What is needed"], [[a, b] for a, b in items], widths=[5.4, 11.4])

doc.add_paragraph()
rule(doc, "DDDDDD", 6)
para(doc, "Grounded in a live read of the platform database on 5 August 2026 — 11 organisations; "
          "Micromart at 162 borrowers, 199 loans, 5 products, 2 workflows, M-Pesa and CRB configured "
          "— together with the decision engine, the ServiceSuite bridge, the console navigation "
          "registry, the Connected Suite definition, the 62-model schema, the Hub app-store seed, and "
          "Micromart’s own Micro Eazy Monthly product screen.",
     9, italic=True, color=MUTED)

os.makedirs("reports", exist_ok=True)
doc.save(OUT)
print("Wrote " + OUT)
